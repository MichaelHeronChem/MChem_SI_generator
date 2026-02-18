import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import os
import math

def format_sig_figs(val, n):
    if pd.isna(val) or val == 0: return "0"
    decimals = n - 1 - int(math.floor(math.log10(abs(val))))
    return format(val, f'.{decimals}f') if decimals > 0 else format(round(val, decimals), '.0f')

class SI_PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 10)
        self.cell(0, 10, 'Supporting Information', 0, 1, 'R')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def print_table_header(self, headers, widths):
        self.set_font('Arial', 'B', 8)
        self.set_fill_color(240, 240, 240) # Light gray background
        for i, h in enumerate(headers):
            self.cell(widths[i], 8, h, 1, 0, 'C', True)
        self.ln()

    def print_row(self, data, widths):
        """
        Custom function to handle text wrapping in the first column (Name).
        Calculates required height based on the text length.
        """
        self.set_font('Arial', '', 8)
        
        # 1. Calculate height needed for the Name (first column)
        x_start = self.get_x()
        y_start = self.get_y()
        
        # Simulate printing the name to see how high it goes
        name_width = widths[0]
        # Calculate number of lines name will take
        # approximate char limit per line ~ width / 2 (heuristic for Arial 8)
        num_lines = self.get_string_width(data[0]) / name_width
        num_lines = math.ceil(num_lines) if num_lines > 1 else 1
        
        # Force a minimum height of 6mm, scale up if wrapped
        row_height = 6 * num_lines 
        
        # 2. Print the First Column (Wrapped)
        self.multi_cell(widths[0], 6, data[0], 1, 'L')
        
        # 3. Print the rest of the columns (Fixed height)
        # Move cursor back to top-right of the first cell
        self.set_xy(x_start + widths[0], y_start)
        
        for i in range(1, len(data)):
            self.cell(widths[i], row_height, str(data[i]), 1, 0, 'C')
        
        self.ln()

def generate_si_files(csv_input, nmr_image_folder, output_base):
    df = pd.read_csv(csv_input)
    doc = Document()
    pdf = SI_PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # --- 1. GLOBAL EXPERIMENTAL METHODS ---
    doc.add_heading('General Experimental Procedures', level=1)
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, 'General Experimental Procedures', 0, 1)
    pdf.ln(5)

    methods = [
        ("Amine Stock Solutions", "Amine stock solutions were made as follows: Amine and HMDSO was added to 10 mL volumetric flask and filled up to the mark with 15% MeCN-d3 in MeCN dried over molecular sieves to give 10 mL of amine stock solution."),
        ("Aldehyde Stock Solutions", "Aldehyde stock solutions were made as follows: Aldehyde was added to a sample vial. Enough MeCN was dispensed into the sample vial by the Hamilton liquid handler using 1000 uL disposable tips to create a 100.00 mM aldehyde solution."),
        ("Reaction Method", "Reactions were made up using Hamilton liquid handler with 300 uL disposable tips. Teflon coated stirrer bars were added to the vials, which were then sealed and left to stir overnight before NMR analysis.")
    ]

    for title, text in methods:
        # Word
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(text)
        # PDF
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 6, title, 0, 1)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, text)
        pdf.ln(3)

    # --- 2. BLOCK DATA ---
    blocks = df.groupby('Experiment_Block')
    
    # Global Table Counter
    tbl_counter = 1

    for block_id, block_data in blocks:
        doc.add_page_break()
        pdf.add_page()
        
        # Block Title
        block_title = f"Imine Formation Reactions - Block {block_id}"
        doc.add_heading(block_title, level=1)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, block_title, 0, 1)
        pdf.ln(5)

        # --- AMINE TABLE ---
        lbl_amine = f"Table {tbl_counter}: Amine Stock Solutions (Block {block_id})"
        doc.add_paragraph(lbl_amine).bold = True
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, lbl_amine, 0, 1)
        tbl_counter += 1

        amines = block_data.drop_duplicates(subset=['Amine_Name'])
        headers = ['Amine Name', 'Mass (mg)', 'Amt (mmol)', 'Conc (mM)', 'HMDSO (mmol)']
        # PDF Widths: Name gets 60mm, others get smaller space
        widths = [60, 25, 25, 25, 30] 
        
        # Word Setup
        t1 = doc.add_table(rows=1, cols=5)
        t1.style = 'Table Grid'
        for i, h in enumerate(headers): t1.rows[0].cells[i].text = h
        
        # PDF Setup
        pdf.print_table_header(headers, widths)

        for _, row in amines.iterrows():
            # CORRECTION: Removed * 1000 multiplier
            m_mg = row['Actual_Mass_Amine_g'] 
            mmol = m_mg / row['Amine_MW_g_mol'] # Assuming mass is now correct for MW calc? 
            # Note: If mass is mg and MW is g/mol, result is mmol. Correct.

            vals = [
                str(row['Amine_Name']), 
                f"{m_mg:.2f}", 
                f"{mmol:.3f}", 
                format_sig_figs(row['Actual_Amine_Conc_mM'], 3), 
                f"{row['Actual_Mass_HMDSO_mg']/162.38:.3f}"
            ]
            
            # Word Row
            cells = t1.add_row().cells
            for i, v in enumerate(vals): cells[i].text = v
            
            # PDF Row (Smart Wrapping)
            pdf.print_row(vals, widths)

        # --- ALDEHYDE TABLE ---
        pdf.ln(5)
        lbl_ald = f"Table {tbl_counter}: Aldehyde Stock Solutions (Block {block_id})"
        doc.add_paragraph(lbl_ald).bold = True
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, lbl_ald, 0, 1)
        tbl_counter += 1

        aldehydes = block_data.drop_duplicates(subset=['Aldehyde_Name'])
        headers = ['Aldehyde Name', 'Mass (mg)', 'Amt (mmol)', 'Vol (uL)', 'Conc (mM)']
        widths = [60, 25, 25, 25, 30]
        
        t2 = doc.add_table(rows=1, cols=5)
        t2.style = 'Table Grid'
        for i, h in enumerate(headers): t2.rows[0].cells[i].text = h
        pdf.print_table_header(headers, widths)

        for _, row in aldehydes.iterrows():
            m_mg = row['Aldehyde_Actual_Mass_mg']
            vals = [
                str(row['Aldehyde_Name']), 
                f"{m_mg:.2f}", 
                f"{m_mg/row['Aldehyde_MW_g_mol']:.3f}", 
                f"{row['Aldehyde_Vol_Required_uL']:.1f}", 
                format_sig_figs(row['Actual_Aldehyde_Conc_mM'], 4)
            ]
            
            cells = t2.add_row().cells
            for i, v in enumerate(vals): cells[i].text = v
            pdf.print_row(vals, widths)

        # --- REACTION TABLE ---
        pdf.ln(5)
        lbl_rxn = f"Table {tbl_counter}: Reaction Volumes (Block {block_id})"
        doc.add_paragraph(lbl_rxn).bold = True
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, lbl_rxn, 0, 1)
        tbl_counter += 1

        headers = ['Amine', 'Aldehyde', 'Vol Amine (uL)', 'Vol Aldehyde (uL)']
        widths = [50, 50, 35, 35]
        
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Table Grid'
        for i, h in enumerate(headers): t3.rows[0].cells[i].text = h
        pdf.print_table_header(headers, widths)

        for _, row in block_data.iterrows():
            vals = [
                str(row['Amine_Name']), 
                str(row['Aldehyde_Name']), 
                f"{row['Vol_Amine_Sol_uL']:.1f}", 
                f"{row['Vol_Aldehyde_Sol_uL']:.1f}"
            ]
            cells = t3.add_row().cells
            for i, v in enumerate(vals): cells[i].text = v
            # Use same smart row printer because reaction names can be long too
            pdf.print_row(vals, widths)

    # --- 3. NMR APPENDIX ---
    doc.add_page_break()
    pdf.add_page()
    doc.add_heading('Appendix: 1H NMR Spectra', level=1)
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, "Appendix: 1H NMR Spectra", 0, 1)

    for _, row in df.iterrows():
        title = f"1H NMR: {row['Amine_Name']} + {row['Aldehyde_Name']}"
        img_name = f"{str(row['Amine_Name']).replace(' ', '_')}_{str(row['Aldehyde_Name']).replace(' ', '_')}.png"
        img_path = os.path.join(nmr_image_folder, img_name)
        
        doc.add_paragraph(title).bold = True
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, title, 0, 1)

        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
            if pdf.get_y() > 180: pdf.add_page()
            pdf.image(img_path, w=170)
            pdf.ln(5)
        else:
            doc.add_paragraph(f"[Image Missing: {img_name}]")
            pdf.cell(0, 10, f"[Image Missing]", 1, 1)

    doc.save(f"{output_base}.docx")
    pdf.output(f"{output_base}.pdf")
    print(f"Success! Generated {output_base}.docx and {output_base}.pdf")

if __name__ == "__main__":
    generate_si_files("data/output/extracted_reaction_data.csv", "data/raw/nmr_images", "data/output/Final_SI_Report")